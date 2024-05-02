from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from core.services.base import BaseDocumentService


class WordService(BaseDocumentService):
    def __init__(self):
        self.document = None

    def load(self, file) -> None:
        self.document = Document(file)

    def update(self, params: dict) -> None:
        if not self.document:
            raise ValueError("Word файл не загружен.")

        # Ваш код для обновления Word-документа
        def find_bookmarks(element, data):
            for bookmark in element.findall('.//w:bookmarkStart', self.document.element.nsmap):
                name = bookmark.get(qn('w:name'))
                if name in data:
                    parent = bookmark.getparent()

                    if isinstance(parent, CT_P):  # Проверка, является ли родительский элемент параграфом
                        if not isinstance(data[name], list):
                            print(f"Закладка {name} уже находится в параграфе.")
                            run = OxmlElement('w:r')
                            t = OxmlElement('w:t')
                            t.text = data[name]
                            run.append(t)
                            parent.append(run)
                        else:
                            table = None
                            for ancestor in parent.iterancestors():  # Перебор всех предков элемента
                                if isinstance(ancestor, CT_Tbl):  # Проверка, является ли элемент таблицей
                                    table = ancestor
                                    break
                            if table is not None:  # Если таблица найдена
                                # Находим соответствующий объект Table в документе
                                for tbl in self.document.tables:
                                    if tbl._element == table:
                                        # Удаляем все строки
                                        for row in tbl.rows[3:]:
                                            tbl._element.remove(row._element)
                                        # Добавляем новые строки
                                        items = data[name]
                                        # Если элемент является одиночным массивом, оборачиваем его в дополнительный массив
                                        if not isinstance(items[0], list):
                                            items = [items]
                                        for item in items:
                                            row = tbl.add_row()
                                            cells = row.cells
                                            for i in range(len(item)):
                                                cells[i].text = item[i]
                                        break

                            else:
                                print(f"Таблица не найдена для закладки {name}")
                    else:
                        print(f"Закладка {name} не находится в параграфе.")
                        new_paragraph = self.document.add_paragraph(data[name])
                        parent.addnext(new_paragraph._element)

        # Перебор всех элементов документа
        for element in self.document.element.body:
            find_bookmarks(element, params)

        # Перебор всех колонтитулов документа
        for section in self.document.sections:
            find_bookmarks(section.header._element, params)
            find_bookmarks(section.footer._element, params)
        # Установка границ для каждой ячейки в таблице
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Получение свойств ячейки
                    tcPr = cell._tc.get_or_add_tcPr()
                    # Создание элемента границы ячейки
                    tcBorders = OxmlElement('w:tcBorders')
                    # Создание и настройка элементов границ
                    for border in ['top', 'left', 'bottom', 'right']:
                        element = OxmlElement(f'w:{border}')
                        element.set(qn('w:val'), 'single')
                        element.set(qn('w:sz'), '4')
                        element.set(qn('w:space'), '0')
                        element.set(qn('w:color'), 'auto')
                        tcBorders.append(element)
                    # Добавление границ в свойства ячейки
                    tcPr.append(tcBorders)
    def save_to_bytes(self) -> BytesIO:
        if not self.document:
            raise ValueError("Word файл не загружен.")
        output = BytesIO()
        self.document.save(output)
        output.seek(0)
        return output

    def save_to_file(self, file_path: str) -> None:
        if self.document:
            self.document.save(file_path)
        else:
            raise ValueError("Word файл не загружен.")
